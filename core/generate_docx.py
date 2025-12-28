from docx import Document


def create_docx(papers, output_file):
    doc = Document()
    doc.add_heading("Research Summary", level=1)

    for i, p in enumerate(papers, 1):
        doc.add_heading(f"{i}. {p.get('title','Unknown')}", level=2)
        doc.add_paragraph(f"Authors: {p.get('authors','')}")
        doc.add_paragraph(f"Year: {p.get('year','')}")
        doc.add_paragraph("Abstract:")
        doc.add_paragraph(p.get("content") or p.get("abstract",""))

        if p.get("pdf_url"):
            doc.add_paragraph(f"PDF: {p['pdf_url']}")
        elif p.get("link"):
            doc.add_paragraph(f"Link: {p['link']}")

    doc.save(output_file)
    return output_file
